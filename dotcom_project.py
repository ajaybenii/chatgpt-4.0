import pandas as pd
from docx import Document
import asyncio

def format_word_document(doc: Document, project_name: str, project_id: str):
    """Format and add content to Word document"""
    # Add project name as heading
    doc.add_heading(f"Project Name: {project_name}", level=2)
    doc.add_heading(f"Project Id: {project_id}", level=2)

    # Add page break
    doc.add_page_break()

async def generate_project_descriptions(csv_path: str = "projects.csv", output_path: str = "project_descriptions.docx"):
    """Main function to process CSV and generate Word document"""
    try:
        # Read CSV file with utf-8 encoding to handle special characters
        df = pd.read_csv(csv_path, encoding='utf-8')
        
        # Create Word document
        doc = Document()
        # Process each project
        for _, row in df.iterrows():
            project_id = str(row['projectId'])
            project_name = str(row['Projectname'])
            project_type = row['Projectype']
            
            print(f"Processing {project_name}...")
            
            try:
                # Ensure the formatting function is correctly called
                format_word_document(doc, project_name, project_id)
                print(f"Successfully processed {project_name}")
                    
            except Exception as e:
                print(f"Error processing {project_name}: {str(e)}")
                doc.add_heading(project_name, level=1)
                doc.add_paragraph(f"Error processing this project: {str(e)}")
                doc.add_page_break()
        
        # Save the document
        doc.save(output_path)
        print(f"Document saved successfully to {output_path}!")
        return True
        
    except Exception as e:
        print(f"Error processing CSV: {str(e)}")
        return False

# Example usage
if __name__ == "__main__":
    asyncio.run(generate_project_descriptions("project_ids.csv", "project_descriptions.docx"))
